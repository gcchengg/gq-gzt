#!/usr/bin/env python3
"""Generate an editable FAW-styled PRD DOCX from confirmed Markdown."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


FAW_BLUE = "152A8C"
FAW_RED = "C00000"
TEXT = "344054"
MUTED = "667085"
LIGHT_BLUE = "EAF0FF"
PREFERRED_FONT = "Microsoft YaHei"
FALLBACK_FONT = "Arial Unicode MS"
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*$")


def available_font():
    mac_font = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
    yahei_candidates = (
        Path("/Library/Fonts/Microsoft YaHei.ttf"),
        Path.home() / "Library/Fonts/Microsoft YaHei.ttf",
    )
    if any(path.exists() for path in yahei_candidates):
        return PREFERRED_FONT
    if mac_font.exists():
        return FALLBACK_FONT
    return PREFERRED_FONT


FONT = available_font()


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_paragraph_bottom_border(paragraph, color=FAW_BLUE, size=12, space=4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35

    for name, size, before, after in (
        ("Heading 1", 16, 14, 8),
        ("Heading 2", 13, 11, 6),
        ("Heading 3", 11.5, 8, 5),
    ):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(FAW_BLUE)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor.from_string(TEXT)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    set_run_font(run, size=8, color=MUTED)


def configure_page(document, logo_path):
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.6))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(11.8)
    table.columns[1].width = Cm(4.8)
    left = table.cell(0, 0).paragraphs[0]
    run = left.add_run("产品需求文档 / PRD")
    set_run_font(run, size=8.5, bold=True, color=FAW_BLUE)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if logo_path.exists():
        right.add_run().add_picture(str(logo_path), width=Cm(3.4))
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run("一汽股权｜需求文档")
    set_run_font(fr, size=8, color=MUTED)
    add_page_number(footer.add_paragraph())


def add_cover(document, title, subtitle, metadata):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PRD / 需求文档")
    set_run_font(r, size=9, bold=True, color=FAW_RED)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color=FAW_BLUE)
    set_paragraph_bottom_border(p, color=FAW_BLUE, size=14, space=7)

    if subtitle:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
        r = p.add_run(subtitle)
        set_run_font(r, size=12, color=MUTED)

    if metadata:
        table = document.add_table(rows=len(metadata), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        table.columns[0].width = Cm(3.2)
        table.columns[1].width = Cm(12.8)
        for i, (key, value) in enumerate(metadata):
            label_cell, value_cell = table.rows[i].cells
            set_cell_shading(label_cell, LIGHT_BLUE)
            set_cell_shading(value_cell, "F7F9FC")
            label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            value_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            lr = label_cell.paragraphs[0].add_run(key)
            vr = value_cell.paragraphs[0].add_run(value)
            set_run_font(lr, size=9, bold=True, color=FAW_BLUE)
            set_run_font(vr, size=9.5, color=TEXT)
            set_cell_margins(label_cell)
            set_cell_margins(value_cell)
        document.add_paragraph()


def split_table_row(line):
    return [x.strip() for x in line.strip().strip("|").split("|")]


def add_markdown_table(document, rows):
    if len(rows) < 2:
        return
    header = split_table_row(rows[0])
    data_lines = rows[2:] if TABLE_SEPARATOR_RE.match(rows[1]) else rows[1:]
    table = document.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for j, value in enumerate(header):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, FAW_BLUE)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        run = cell.paragraphs[0].add_run(value)
        set_run_font(run, size=9, bold=True, color="FFFFFF")
        set_cell_margins(cell, top=120, bottom=120)
    set_repeat_table_header(table.rows[0])
    for i, line in enumerate(data_lines):
        values = split_table_row(line)
        cells = table.add_row().cells
        for j in range(len(header)):
            value = values[j] if j < len(values) else ""
            cell = cells[j]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i % 2:
                set_cell_shading(cell, "F7F9FC")
            cell.text = ""
            run = cell.paragraphs[0].add_run(value)
            set_run_font(run, size=9, color=TEXT)
            set_cell_margins(cell, top=105, bottom=105)
    document.add_paragraph()


def resolve_image(markdown_path, raw_path):
    value = raw_path.replace("%20", " ").strip()
    candidate = Path(value)
    if not candidate.is_absolute():
        candidate = (markdown_path.parent / candidate).resolve()
    return candidate


def add_image(document, image_path, alt_text):
    if not image_path.exists():
        raise FileNotFoundError(f"Markdown image not found: {image_path}")
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Cm(15.8))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    run = caption.add_run(alt_text or image_path.name)
    set_run_font(run, size=8, color=MUTED)


def parse_metadata(lines):
    metadata = []
    for line in lines:
        if not line.startswith(">"):
            continue
        content = line[1:].strip().replace("  ", "")
        if "：" in content:
            key, value = content.split("：", 1)
            metadata.append((key.strip(), value.strip()))
    return metadata


def generate(input_md: Path | str, output_docx: Path | str) -> Path:
    input_md = Path(input_md).resolve()
    output_docx = Path(output_docx).resolve()
    lines = input_md.read_text(encoding="utf-8").splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "产品需求文档")
    metadata = parse_metadata(lines)
    skill_root = Path(__file__).resolve().parent.parent
    logo = skill_root / "assets/icons/faw-equity-logo.png"

    document = Document()
    configure_styles(document)
    configure_page(document, logo)
    add_cover(document, title, "需求变更与页面说明", metadata)

    table_buffer = []
    skip_title = True
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("|"):
            table_buffer.append(line)
            i += 1
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table_buffer.append(lines[i].rstrip())
                i += 1
            add_markdown_table(document, table_buffer)
            table_buffer = []
            continue
        if line.startswith("# "):
            if skip_title:
                skip_title = False
            else:
                document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            p = document.add_heading(line[3:].strip(), level=1)
            set_paragraph_bottom_border(p, color=FAW_BLUE, size=8, space=4)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith(">"):
            pass
        elif IMAGE_RE.fullmatch(line.strip()):
            match = IMAGE_RE.fullmatch(line.strip())
            add_image(document, resolve_image(input_md, match.group(2)), match.group(1))
        elif re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = document.add_paragraph(style="List Bullet")
            run = p.add_run(text)
            set_run_font(run, size=10.5, color=TEXT)
        elif re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = document.add_paragraph(style="List Number")
            run = p.add_run(text)
            set_run_font(run, size=10.5, color=TEXT)
        elif line.strip():
            p = document.add_paragraph()
            run = p.add_run(line.strip())
            set_run_font(run, size=10.5, color=TEXT)
        i += 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.core_properties.title = title
    document.core_properties.subject = "产品需求文档"
    document.core_properties.author = "一汽股权"
    document.save(output_docx)
    return output_docx


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    print(generate(args.input, args.output))


if __name__ == "__main__":
    main()
