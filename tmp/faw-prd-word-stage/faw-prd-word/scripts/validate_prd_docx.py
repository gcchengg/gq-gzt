#!/usr/bin/env python3
"""Validate the structural requirements of a generated FAW PRD DOCX."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

from docx import Document


IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
UNRESOLVED = ("TODO", "TBD", "{{", "}}", "[[", "]]")


def validate(docx_path: Path | str, source_markdown: Path | str | None = None):
    docx_path = Path(docx_path).resolve()
    if not docx_path.exists() or docx_path.stat().st_size == 0:
        raise ValueError(f"DOCX does not exist or is empty: {docx_path}")

    document = Document(docx_path)
    text = "\n".join(p.text for p in document.paragraphs)
    headings = sum(1 for p in document.paragraphs if p.style.name.startswith("Heading"))
    if headings < 1:
        raise ValueError("DOCX contains no real Word headings")
    for marker in UNRESOLVED:
        if marker in text:
            raise ValueError(f"Unresolved marker found: {marker}")

    with zipfile.ZipFile(docx_path) as archive:
        names = archive.namelist()
        embedded_images = len([n for n in names if n.startswith("word/media/")])
        combined_xml = (
            archive.read("word/document.xml").decode("utf-8")
            + archive.read("word/styles.xml").decode("utf-8")
        )
    for color in ("152A8C", "C00000"):
        if color not in combined_xml:
            raise ValueError(f"Required brand color missing: #{color}")

    source_images = 0
    if source_markdown:
        source = Path(source_markdown).resolve()
        markdown = source.read_text(encoding="utf-8")
        source_images = len(IMAGE_RE.findall(markdown))
        if embedded_images < source_images + 1:
            raise ValueError(
                f"Expected logo plus {source_images} source image(s), found {embedded_images}"
            )

    return {
        "valid": True,
        "headings": headings,
        "embedded_images": embedded_images,
        "source_images": source_images,
        "bytes": docx_path.stat().st_size,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--source-markdown", type=Path)
    args = parser.parse_args()
    print(json.dumps(validate(args.docx, args.source_markdown), ensure_ascii=False))


if __name__ == "__main__":
    main()
